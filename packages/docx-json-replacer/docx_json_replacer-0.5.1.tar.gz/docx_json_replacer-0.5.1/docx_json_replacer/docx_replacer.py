"""
DocxReplacer with table support for docx-json-replacer
"""
import json
import re
from typing import Dict, Any, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from .utility.html_parse import clean_html_content
    from .table_handler import TableHandler
except ImportError:
    from utility.html_parse import clean_html_content
    from table_handler import TableHandler


class DocxReplacer:
    """Replace placeholders in DOCX files with JSON data, including table support"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.table_handler = TableHandler()
        self.table_placeholders = {}
        self._value_cache = {}
        self._table_check_cache = {}

    def replace_from_json(self, json_data: Dict[str, Any]) -> None:
        """Replace placeholders in paragraphs AND tables"""

        # Pre-compile patterns
        patterns = self._compile_patterns(json_data)

        # Pre-process values
        processed_values = self._preprocess_values(json_data)

        # Process regular paragraphs
        self._process_paragraphs(patterns, processed_values)

        # Process table cells
        self._process_tables(patterns, processed_values)

        # Insert dynamic tables for table data
        self._batch_insert_tables()

    def _compile_patterns(self, json_data: Dict[str, Any]) -> Dict[str, Tuple[re.Pattern, re.Pattern]]:
        """Pre-compile regex patterns for all placeholders"""
        patterns = {}
        for key in json_data.keys():
            escaped_key = re.escape(key)
            pattern = re.compile(r'\{\{' + escaped_key + r'\}\}')
            pattern_spaced = re.compile(r'\{\{ ' + escaped_key + r' \}\}')
            patterns[key] = (pattern, pattern_spaced)
        return patterns

    def _preprocess_values(self, json_data: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        """Pre-process all values"""
        processed = {}

        for key, value in json_data.items():
            is_table = self._check_is_table(value)

            if is_table:
                processed[key] = {
                    'is_table': True,
                    'original': value,
                    'processed': None
                }
            else:
                cleaned = clean_html_content(value) if isinstance(value, str) else str(value)
                processed[key] = {
                    'is_table': False,
                    'original': value,
                    'processed': cleaned
                }

        return processed

    def _check_is_table(self, value: Any) -> bool:
        """Check if value is table data with caching"""
        value_id = id(value)
        if value_id in self._table_check_cache:
            return self._table_check_cache[value_id]

        result = (self.table_handler.is_table_data(value) or
                  (isinstance(value, str) and '<table' in value.lower()))

        self._table_check_cache[value_id] = result
        return result

    def _process_paragraphs(self, patterns: Dict, processed_values: Dict) -> None:
        """Process all document paragraphs"""
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            if not text or '{{' not in text:
                continue

            modified = False
            new_text = text

            for key, (pattern, pattern_spaced) in patterns.items():
                if pattern.search(new_text) or pattern_spaced.search(new_text):
                    value_data = processed_values[key]

                    if value_data['is_table']:
                        # Store for table insertion
                        self.table_placeholders[paragraph] = (key, value_data['original'])
                        new_text = pattern.sub('', new_text)
                        new_text = pattern_spaced.sub('', new_text)
                        modified = True
                    else:
                        # Regular text replacement
                        replacement = value_data['processed']
                        new_text = pattern.sub(replacement, new_text)
                        new_text = pattern_spaced.sub(replacement, new_text)
                        modified = True

            if modified:
                paragraph.text = new_text

    def _process_tables(self, patterns: Dict, processed_values: Dict) -> None:
        """Process all table cells in the document"""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Process each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if not text or '{{' not in text:
                            continue

                        modified = False
                        new_text = text

                        for key, (pattern, pattern_spaced) in patterns.items():
                            if pattern.search(new_text) or pattern_spaced.search(new_text):
                                value_data = processed_values[key]

                                if value_data['is_table']:
                                    # For table data in cells, just clear the placeholder
                                    # (inserting tables inside cells is complex)
                                    new_text = pattern.sub('[Table data - see document]', new_text)
                                    new_text = pattern_spaced.sub('[Table data - see document]', new_text)
                                    modified = True
                                else:
                                    # Regular text replacement
                                    replacement = value_data['processed']
                                    new_text = pattern.sub(replacement, new_text)
                                    new_text = pattern_spaced.sub(replacement, new_text)
                                    modified = True

                        if modified:
                            paragraph.text = new_text

    def _batch_insert_tables(self) -> None:
        """Insert all tables in batch"""
        if not self.table_placeholders:
            return

        processed_tables = {}
        for paragraph, (key, value) in self.table_placeholders.items():
            table_data = self.table_handler.process_table_data(value)
            if table_data.get('rows'):
                processed_tables[paragraph] = table_data

        for paragraph, table_data in processed_tables.items():
            self._insert_table_fast(paragraph, table_data)

    def _insert_table_fast(self, paragraph, table_data: Dict[str, Any]) -> None:
        """Fast table insertion"""
        rows = table_data['rows']
        num_rows = len(rows)
        num_cols = len(rows[0]['cells']) if rows and 'cells' in rows[0] else 0

        if num_rows == 0 or num_cols == 0:
            return

        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        table_rows = table.rows
        for row_idx, row_data in enumerate(rows):
            cells = row_data.get('cells', [])
            style = row_data.get('style', {})
            row_cells = table_rows[row_idx].cells

            for col_idx, cell_text in enumerate(cells):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]
                    cell.text = str(cell_text)

                    if style:
                        self._apply_cell_style_fast(cell, style)

        parent.insert(index + 1, table._element)

    def _apply_cell_style_fast(self, cell, style: Dict[str, Any]) -> None:
        """Fast cell styling"""
        if bg := style.get('bg'):
            self._set_cell_bg_fast(cell, bg)

        if any(style.get(k) for k in ['bold', 'italic', 'color']):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if style.get('bold'):
                        run.bold = True
                    if style.get('italic'):
                        run.italic = True

                    if color_hex := style.get('color'):
                        color_hex = color_hex.replace('#', '')
                        if len(color_hex) == 6:
                            rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                            run.font.color.rgb = RGBColor(*rgb)

    def _set_cell_bg_fast(self, cell, color: str) -> None:
        """Fast background setting"""
        color = color.replace('#', '')
        tc_pr = cell._tc.get_or_add_tcPr()

        if existing := tc_pr.find(qn('w:shd')):
            tc_pr.remove(existing)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tc_pr.append(shd)

    def save(self, output_path: str) -> None:
        """Save the document"""
        self.doc.save(output_path)

    def replace_from_json_file(self, json_path: str) -> None:
        """Load JSON and replace"""
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        self.replace_from_json(json_data)


def replace_docx_template(docx_path: str, json_data: Dict[str, Any], output_path: str) -> None:
    """Utility function to replace template and save in one step"""
    replacer = DocxReplacer(docx_path)
    replacer.replace_from_json(json_data)
    replacer.save(output_path)