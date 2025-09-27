# import pandas as pd
# import markdown
# from docx import Document
# from fpdf import FPDF

# def make_codebook(df, output_format='markdown', filename='codebook'):
#     """
#     Generate a codebook from a pandas DataFrame and export it in the specified format.
    
#     Parameters:
#     - df: pandas DataFrame
#     - output_format: 'markdown', 'docx', or 'pdf'
#     - filename: base name for the output file
#     """

#     # Build a codebook as a DataFrame
#     codebook_df = pd.DataFrame({
#         'Variable': df.columns,
#         'Data Type': [df[col].dtype for col in df.columns],
#         'Non-Null Count': [df[col].count() for col in df.columns],
#         'Unique Values': [df[col].nunique() for col in df.columns],
#         'Sample Values': [df[col].dropna().unique()[:3] for col in df.columns],
#         'Description': ['' for _ in df.columns]  # Placeholder, you can customize
#     })

#     # MARKDOWN output
#     if output_format == 'markdown':
#         md = codebook_df.to_markdown(index=False)
#         with open(f'{filename}.md', 'w') as f:
#             f.write(f'# Codebook\n\n{md}')

#     # DOCX output
#     elif output_format == 'docx':
#         doc = Document()
#         doc.add_heading('Codebook', level=1)
#         table = doc.add_table(rows=1, cols=len(codebook_df.columns))
#         hdr_cells = table.rows[0].cells
#         for i, col in enumerate(codebook_df.columns):
#             hdr_cells[i].text = col

#         for _, row in codebook_df.iterrows():
#             row_cells = table.add_row().cells
#             for i, item in enumerate(row):
#                 row_cells[i].text = str(item)

#         doc.save(f'{filename}.docx')

#     # PDF output
#     elif output_format == 'pdf':
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=10)
#         pdf.cell(200, 10, txt="Codebook", ln=True, align='C')

#         for _, row in codebook_df.iterrows():
#             line = ', '.join([f"{k}: {v}" for k, v in row.items()])
#             pdf.multi_cell(0, 10, line)

#         pdf.output(f"{filename}.pdf")

#     else:
#         raise ValueError("Invalid format. Choose 'markdown', 'docx', or 'pdf'.")

#     print(f"Codebook saved as {filename}.{output_format}")

# --------------------

import pandas as pd
import numpy as np
from docx import Document
from fpdf import FPDF
import textwrap

def make_codebook(df, output_format='markdown', filename='codebook', custom_descriptions=None):
    """
    Generate a skimr-style codebook from a pandas DataFrame and export it.
    
    Parameters:
    - df: pandas DataFrame
    - output_format: 'markdown', 'docx', or 'pdf'
    - filename: base name for output file
    - custom_descriptions: dict of {column_name: description} overrides
    """

    def describe_variable(col, series):
        if pd.api.types.is_numeric_dtype(series):
            return {
                'n': len(series),
                'missing': series.isna().sum(),
                'mean': round(series.mean(), 2),
                'std': round(series.std(), 2),
                'min': series.min(),
                'max': series.max()
            }
        elif pd.api.types.is_categorical_dtype(series) or series.dtype == object:
            top = series.mode().iloc[0] if not series.mode().empty else ''
            return {
                'n': len(series),
                'missing': series.isna().sum(),
                'unique': series.nunique(),
                'top': top,
                'freq': series.value_counts().iloc[0] if not series.value_counts().empty else ''
            }
        else:
            return {
                'n': len(series),
                'missing': series.isna().sum(),
                'example': str(series.dropna().iloc[0]) if not series.dropna().empty else ''
            }

    rows = []
    for col in df.columns:
        desc = custom_descriptions.get(col) if custom_descriptions else None
        if not desc:
            desc = col.replace('_', ' ').capitalize()
        summary = describe_variable(col, df[col])
        summary['variable'] = col
        summary['type'] = str(df[col].dtype)
        summary['description'] = desc
        rows.append(summary)

    codebook_df = pd.DataFrame(rows)

    # Reorder columns
    cols_order = ['variable', 'type', 'description'] + [col for col in codebook_df.columns if col not in ['variable', 'type', 'description']]
    codebook_df = codebook_df[cols_order]

    # MARKDOWN
    if output_format == 'markdown':
        with open(f'{filename}.md', 'w') as f:
            f.write('# Codebook\n\n')
            for _, row in codebook_df.iterrows():
                f.write(f"### {row['variable']} ({row['type']})\n")
                f.write(f"**Description**: {row['description']}\n\n")
                for key, val in row.items():
                    if key not in ['variable', 'type', 'description']:
                        f.write(f"- {key}: {val}\n")
                f.write('\n')

    # DOCX
    elif output_format == 'docx':
        doc = Document()
        doc.add_heading('Codebook', level=1)
        for _, row in codebook_df.iterrows():
            doc.add_heading(f"{row['variable']} ({row['type']})", level=2)
            doc.add_paragraph(f"Description: {row['description']}")
            for key, val in row.items():
                if key not in ['variable', 'type', 'description']:
                    doc.add_paragraph(f"{key}: {val}")
        doc.save(f"{filename}.docx")

    # PDF
    elif output_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.multi_cell(0, 10, txt="Codebook", align='L')
        for _, row in codebook_df.iterrows():
            pdf.set_font("Arial", style='B', size=10)
            pdf.cell(0, 10, f"{row['variable']} ({row['type']})", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 10, f"Description: {row['description']}")
            for key, val in row.items():
                if key not in ['variable', 'type', 'description']:
                    line = f"{key}: {val}"
                    pdf.multi_cell(0, 10, line)
            pdf.ln(5)
        pdf.output(f"{filename}.pdf")

    else:
        raise ValueError("Format must be one of: 'markdown', 'docx', 'pdf'")

    print(f"Codebook saved to {filename}.{output_format}")


df = pd.DataFrame({
    'age': [25, 32, 47, None],
    'gender': ['M', 'F', 'M', 'F'],
    'score': [88.5, 92.0, 85.0, None],
    'notes': ['A', 'B', None, 'C']
})

custom_desc = {
    'age': 'Age of participant in years',
    'gender': 'Self-reported gender identity',
    'score': 'Score on standardized test',
}

make_codebook(df, output_format='docx', filename='skim_style_codebook', custom_descriptions=custom_desc)
